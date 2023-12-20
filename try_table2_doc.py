from docx import Document
from docxcompose.composer import Composer

doc = Document(r"C:\files\file11.docx")
t = doc.tables[0]
cell = t.cell(0, 0)
row_count = len(t.rows)
print(row_count) # длина таблицы
row = t.add_row()
row.cells[0].text = "31.08.2023"
row.cells[1].text = "Отсутствует"
doc.save(r"C:\files\file12.docx")

# master = Document(r"C:\files\anket1.docx")
# composer = Composer(master)
# # doc1 is the docx file getting copied
# doc1 = Document(r"C:\files\file10.docx")
# composer.append(doc1)
# composer.save(r"C:\files\anket_split.docx")