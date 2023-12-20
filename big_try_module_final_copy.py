# -*- coding: utf-8 -*-
"""
Created on Fri May 26 16:01:53 2023

@author: S.Stycenkov
"""

from tkinter import *
from tkinter import messagebox , scrolledtext  
from datetime import datetime
import pandas as pd
from docx import Document
from docxcompose.composer import Composer

#функция исключения лишних и случайных символов в дате
def onlynum(text):
    a = "0123456789"
    txt = ""
    for i in text:
        if i in a:
            txt += i
    return txt

#удаление заполненных полей
def erase_day():
    # a_entry.select_clear()
    a_entry.delete(0,3)
   
def erase_month():
    b_entry.delete(0,3)

def erase_year():
    c_entry.delete(0,6)


def my_txt(text):
    txt = ""
    for i in range(len(text)):
        if text[i] == "/":
            break
        if text[i] == '"':
            txt += ""
        else:
            txt += text[i]
    return txt

def type_format_client(text):
    for i in range(len(text)):
        if text[i] == "/":
            return "F"
            break
    return "U"
    
   
def excel_read_result(maindata):
    file = "C:\\Users\\S.Stycenkov\\Documents\\result\\result.xls"
    xl = pd.read_excel(file)
    longname = ""
    root.geometry('600x600+{}+{}'.format(wm,wh))
    txt=scrolledtext.ScrolledText(width=51, height=28, wrap=WORD, font=("Courier New", 9))  
    txt.place(x = 40, y = 150)
    
    for i in range(len(xl['Участник'])):
        str_xl = xl['Участник'][i] #scrolledtext.Scrolled / Text
        #вставка модулей
        #ФИО клиента в CAPS:
        nameclient = my_txt(str_xl).rstrip(" ").upper()
        UF_type = type_format_client(str_xl)
        
        if UF_type == "F":
            adressname = "C:\\08 ПОДФТ\\анкеты\\ФЛ"
        else:
            adressname = "C:\\08 ПОДФТ\\анкеты\ЮЛ"
        
        # #вставка даты в 1ую таблицу        
        doc = Document(adressname + "\\" + nameclient + "\\table1.docx")
        t = doc.tables[0]
        cell = t.cell(0, 0)
        row_count = len(t.rows)
        # print(row_count) # длина таблицы
        row = t.add_row()
        row.cells[0].text = str(maindata)
        row.cells[1].text = "Отсутствует"
        doc.save(adressname + "\\" + nameclient + "\\table1.docx")
        
        # #вставка даты во 2ую таблицу
        doc2 = Document(adressname + "\\" + nameclient + "\\table2.docx")
        t = doc2.tables[0]
        cell = t.cell(0, 0)
        row_count = len(t.rows)
        # print(row_count) # длина таблицы
        row = t.add_row()
        row.cells[0].text = str(maindata)
        row.cells[1].text = "Амирова Л.Г."
        row.cells[2].text = "Руководитель отдела финансового мониторинга"
        doc2.save(adressname + "\\" + nameclient + "\\table2.docx")
        
        # #объединение файлов
        master = Document(adressname + "\\" + nameclient + "\\anketa.docx")
        composer = Composer(master)
        # doc1 is the docx file getting copied
        # doc = Document(r"C:\files\file8.docx")
        composer.append(doc)
        composer.append(doc2)
        composer.save(adressname + "\\" + nameclient + "\\" + nameclient + "_" + str(maindata) + ".docx")
        
        longname = str(i+1) + ")\t" + "[" + UF_type + "] " + nameclient + "\n"
        txt.insert(END, longname)
        print("[" + UF_type + "] " + my_txt(str_xl).rstrip(" "))
   

def display_full_name():
    
    mday = onlynum(str(a.get()))
    if mday == "" or int(mday) <= 0 or int(mday)>31:
        if mday == "":
            txt1 = "<ничего>"
        else:
            txt1 = mday
        messagebox.showerror(title="Неверная дата", message="Вы ввели " + txt1 + " в дату, будьте внимательны!")
        flag_a = 0
    else:
        flag_a = 1
    
    mmonth = onlynum(str(b.get()))
    if mmonth == "" or int(mmonth) <= 0 or int(mmonth)>12:
        if mmonth == "":
            txt2 = "<ничего>"
        else:
            txt2 = mmonth
        messagebox.showerror(title="Неверная дата", message="Вы ввели " + txt2 + " в месяц, будьте внимательны!")
        flag_b = 0
    else:
        flag_b = 1
    
  
    myear = onlynum(str(c.get()))
    if myear == "" or int(myear) < 2000 or int(myear)>2100:
        if myear == "":
            txt3 = "<ничего>"
        else:
            txt3 = myear
        messagebox.showerror(title="Неверная дата", message="Вы ввели " + txt3 + " в месяц, будьте внимательны!")
        flag_c = 0
    else:
        flag_c = 1  
    
    if (int(mmonth)==2 and int(mday)>29) or ((int(mmonth)==4 or int(mmonth)== 6 or int(mmonth)==9 or int(mmonth)==11) and int(mday)==31):
        flag_a = 0
        txtdata = mdata = nonzero(mday)  + "." + nonzero(mmonth) + "." + myear
        messagebox.showerror(title="Неверная дата", message="Дата " + txtdata  + " ? Вы серьезно? Посмотрите календарь!")    
    
    if (int(mmonth)==2 and int(mday)==29 and int(myear)%4!=0):
        flag_a = 0
        txtdata = mdata = nonzero(mday)  + "." + nonzero(mmonth) + "." + myear
        messagebox.showerror(title="Неверная дата", message= txtdata  + " : 29 февраля в невисокосный год? Really??")    
    
    if flag_a==flag_b==flag_c and flag_a == 1:
        mdata = nonzero(mday)  + "." + nonzero(mmonth) + "." + myear
        answer = messagebox.askyesno(title="Проверка даты:",
                        message="Обновить данные за " + mdata + "? ")
        if answer:
            print(mdata)
            message_button["state"] = DISABLED
            excel_read_result(mdata)


def nonzero(text):
    if int(text)<10:
        return '0' + str(int(text))
    else:
        return str(text)
    
root = Tk()
wm = int(0.5*root.winfo_screenwidth())-210
wh = int(0.5*root.winfo_screenheight())-225
print(root.geometry('600x200+{}+{}'.format(wm,wh)))
 
current_datetime = datetime.now()
cday = current_datetime.day
cmonth = current_datetime.month
cyear = current_datetime.year

root.title("Введите дату для обновления таблиц")

a_label = Label(text="День:", font=("Courier New", 13, 'bold'))
a_label.place(x = 40, y = 50) 
a = StringVar()
a_entry = Entry(textvariable=a, font=("Courier New", 13), width=3, justify=CENTER)
a_entry.place(x = 95, y = 50) 
a_entry.focus()
a_entry.insert(0,nonzero(cday))

b_label = Label(text="Месяц:", font=("Courier New", 13, 'bold'))
b_label.place(x = 150, y = 50) 
b = StringVar()
b_entry = Entry(textvariable=b, font=("Courier New", 13), width=3, justify=CENTER)
b_entry.place(x = 215, y = 50) 
b_entry.insert(0,nonzero(cmonth))


c_label = Label(text="Год:", font=("Courier New", 13, 'bold'))
c_label.place(x = 270, y = 50) 
c = StringVar()
c_entry = Entry(textvariable=c, font=("Courier New", 13), width=6, justify=CENTER)
c_entry.place(x = 315, y = 50) 
c_entry.insert(0,cyear)


message_button = Button(text="обновить\nанкеты", command=display_full_name, font=("Courier New", 13, 'bold'), height=4, width = 12, bg='#B3E5FC', activebackground="white")
message_button.place(x = 424, y = 45)


erase_day_txt = Button(text="очистить\nдни", command=erase_day, font=("Courier New", 11), width=9)
erase_day_txt.place(x = 40, y = 90)

erase_month_txt = Button(text="очистить\nмесяц", command=erase_month, font=("Courier New", 11), width=9)
erase_month_txt.place(x = 155, y = 90)


erase_year_txt = Button(text="очистить\nгод", command=erase_year, font=("Courier New", 11))
erase_year_txt.place(x = 275, y = 90)


root.mainloop()